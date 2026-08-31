from io import BytesIO

import pytest
from docx import Document as DocxDocument

from docmind.extraction.exceptions import CorruptDocumentError
from docmind.extraction.word import extract_docx
from tests.helpers import build_docx


def test_extract_docx_paragraphs() -> None:
    data = build_docx(["First paragraph.", "Second paragraph."])

    result = extract_docx(data)

    assert "First paragraph." in result.text
    assert "Second paragraph." in result.text
    assert result.text.index("First") < result.text.index("Second")
    assert result.tables == []


def test_extract_docx_table() -> None:
    data = build_docx(
        ["Invoice summary"],
        table_rows=[["Item", "Qty"], ["Widget", "3"]],
    )

    result = extract_docx(data)

    assert len(result.tables) == 1
    assert result.tables[0].rows == [["Item", "Qty"], ["Widget", "3"]]
    # Table content should also appear in the flattened text.
    assert "Widget" in result.text


def test_extract_docx_reading_order_preserved() -> None:
    doc = DocxDocument()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Inside table"
    doc.add_paragraph("After table")
    buffer = BytesIO()
    doc.save(buffer)

    result = extract_docx(buffer.getvalue())

    before_idx = result.text.index("Before table")
    table_idx = result.text.index("Inside table")
    after_idx = result.text.index("After table")
    assert before_idx < table_idx < after_idx


def test_extract_docx_corrupt_bytes_raises() -> None:
    with pytest.raises(CorruptDocumentError):
        extract_docx(b"not a real docx file")


def test_extract_docx_detects_headings() -> None:
    doc = DocxDocument()
    doc.add_paragraph("Executive Summary", style="Title")
    doc.add_paragraph("This report covers quarterly results.")
    doc.add_paragraph("Findings", style="Heading 1")
    doc.add_paragraph("Revenue grew steadily this quarter.")
    doc.add_paragraph("Details", style="Heading 2")
    doc.add_paragraph("Broken down by region below.")
    buffer = BytesIO()
    doc.save(buffer)

    result = extract_docx(buffer.getvalue())

    assert [(h.text, h.level) for h in result.headings] == [
        ("Executive Summary", 0),
        ("Findings", 1),
        ("Details", 2),
    ]
    # Each heading's offset should point at exactly where its text starts.
    for heading in result.headings:
        assert result.text[heading.char_offset :].startswith(heading.text)


def test_extract_docx_no_headings_when_no_heading_styles_used() -> None:
    data = build_docx(["Just a normal paragraph.", "Another normal one."])

    result = extract_docx(data)

    assert result.headings == []


def test_extract_docx_reads_document_title_property() -> None:
    doc = DocxDocument()
    doc.core_properties.title = "Quarterly Vendor Report"
    doc.add_paragraph("Body text.")
    buffer = BytesIO()
    doc.save(buffer)

    result = extract_docx(buffer.getvalue())

    assert result.metadata["title"] == "Quarterly Vendor Report"
