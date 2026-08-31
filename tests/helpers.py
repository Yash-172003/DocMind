"""Shared test-fixture builders for real (not mocked) document bytes.

Building real PDF/Word/Excel files via the same libraries the app uses
to read them means these tests exercise real parsing logic end-to-end,
not a mocked stand-in for it.
"""

from io import BytesIO

import pymupdf
from docx import Document as DocxDocument
from openpyxl import Workbook


def build_pdf(pages_text: list[str]) -> bytes:
    doc = pymupdf.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    data: bytes = doc.tobytes()
    doc.close()
    return data


def build_docx(
    paragraphs: list[str], table_rows: list[list[str]] | None = None
) -> bytes:
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_xlsx(sheets: dict[str, list[list[object]]]) -> bytes:
    workbook = Workbook()
    assert workbook.active is not None
    workbook.remove(workbook.active)
    for name, rows in sheets.items():
        sheet = workbook.create_sheet(title=name)
        for row in rows:
            sheet.append(row)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
