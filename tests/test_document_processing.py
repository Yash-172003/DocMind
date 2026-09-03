"""End-to-end tests: upload real document bytes through the API and
verify the background extraction pipeline actually ran, not just that
the endpoints respond.
"""

from collections.abc import AsyncGenerator

import pymupdf
import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy import select

from docmind.core.config import settings
from docmind.db.base import async_session_factory
from docmind.db.models import Chunk
from docmind.main import app
from tests.helpers import build_docx, build_pdf, build_xlsx


@pytest_asyncio.fixture
async def client() -> AsyncGenerator[AsyncClient, None]:
    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://test",
        headers={"X-API-Key": settings.api_key},
    ) as ac:
        yield ac


@pytest.mark.asyncio
async def test_pdf_upload_is_extracted(client: AsyncClient) -> None:
    data = build_pdf(["Invoice total: 1500"])
    files = {"file": ("invoice.pdf", data, "application/pdf")}

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "done"
    assert "Invoice total: 1500" in body["content"]
    assert body["metadata_"]["page_count"] == 1
    assert body["chunk_count"] >= 1


@pytest.mark.asyncio
async def test_pdf_extraction_warnings_survive_to_metadata(
    client: AsyncClient,
) -> None:
    # A genuinely blank page (no text inserted) produces an extraction
    # warning — verifies the pipeline actually carries `warnings` through
    # to the persisted document, not just that extraction detects them.
    doc = pymupdf.open()
    doc.new_page()
    data: bytes = doc.tobytes()
    doc.close()

    files = {"file": ("blank.pdf", data, "application/pdf")}
    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "done"
    assert len(body["metadata_"]["warnings"]) == 1
    assert "no extractable text" in body["metadata_"]["warnings"][0]


@pytest.mark.asyncio
async def test_docx_upload_is_extracted(client: AsyncClient) -> None:
    data = build_docx(["Quarterly report summary"])
    files = {
        "file": (
            "report.docx",
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "done"
    assert "Quarterly report summary" in body["content"]


@pytest.mark.asyncio
async def test_xlsx_upload_is_extracted(client: AsyncClient) -> None:
    data = build_xlsx({"Vendors": [["Name", "Total"], ["Acme Corp", "2000"]]})
    files = {
        "file": (
            "vendors.xlsx",
            data,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    }

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "done"
    assert "Acme Corp" in body["content"]
    assert body["metadata_"]["sheet_count"] == 1
    assert body["metadata_"]["tables"] == [
        {"page_number": 1, "rows": [["Name", "Total"], ["Acme Corp", "2000"]]}
    ]


@pytest.mark.asyncio
async def test_unsupported_file_type_marks_failed_not_crash(
    client: AsyncClient,
) -> None:
    files = {"file": ("archive.zip", b"PK\x03\x04fake", "application/zip")}

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "failed"
    assert body["error_message"] is not None


@pytest.mark.asyncio
async def test_corrupt_pdf_bytes_marks_failed_not_crash(client: AsyncClient) -> None:
    files = {"file": ("broken.pdf", b"not actually a pdf", "application/pdf")}

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    body = content_res.json()

    assert body["status"] == "failed"
    assert body["error_message"] is not None


@pytest.mark.asyncio
async def test_upload_produces_real_persisted_chunks(client: AsyncClient) -> None:
    text = "\n\n".join(
        f"Paragraph number {i} discusses topic {i} in reasonable detail."
        for i in range(20)
    )
    files = {"file": ("report.txt", text.encode(), "text/plain")}

    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    content_res = await client.get(f"/api/v1/documents/{doc_id}/content")
    doc_body = content_res.json()

    chunks_res = await client.get(f"/api/v1/documents/{doc_id}/chunks")
    chunks = chunks_res.json()

    assert doc_body["chunk_count"] == len(chunks)
    assert len(chunks) >= 1
    assert [c["chunk_index"] for c in chunks] == list(range(len(chunks)))
    assert all(c["token_count"] > 0 for c in chunks)
    # Every chunk should be real text pulled from the document, not empty.
    assert all(c["text"].strip() for c in chunks)


@pytest.mark.asyncio
async def test_upload_populates_real_embeddings(client: AsyncClient) -> None:
    # Not exposed via the API (a 1024-float array isn't useful in JSON),
    # so this checks the actual persisted column via a direct DB query —
    # verifying the pipeline wiring, not just that Embedder works in
    # isolation (already covered in tests/embedding/).
    files = {"file": ("note.txt", b"A short note to embed.", "text/plain")}
    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]

    await client.get(f"/api/v1/documents/{doc_id}/content")  # wait for processing

    async with async_session_factory() as db:
        result = await db.execute(
            select(Chunk).where(Chunk.document_id == doc_id)
        )
        chunks = result.scalars().all()

    assert len(chunks) >= 1
    for chunk in chunks:
        assert chunk.embedding is not None
        # 1024 dims: BAAI/bge-large-en-v1.5's real output size, matching
        # the Vector(1024) column set up back in Week 5-6.
        assert len(chunk.embedding) == 1024


@pytest.mark.asyncio
async def test_upload_persists_section_heading_from_real_docx_headings(
    client: AsyncClient,
) -> None:
    from io import BytesIO

    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Findings", style="Heading 1")
    doc.add_paragraph("Revenue grew steadily this quarter across all regions.")
    buffer = BytesIO()
    doc.save(buffer)

    files = {
        "file": (
            "report.docx",
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    upload_res = await client.post("/api/v1/documents/upload", files=files)
    doc_id = upload_res.json()["id"]
    await client.get(f"/api/v1/documents/{doc_id}/content")

    chunks_res = await client.get(f"/api/v1/documents/{doc_id}/chunks")
    chunks = chunks_res.json()

    assert chunks[0]["section_heading"] == "Findings"
