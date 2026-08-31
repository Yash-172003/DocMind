from io import BytesIO

from docx import Document as DocxDocument

from docmind.chunking.structural import chunk_structural
from docmind.extraction.models import ExtractedPage, ExtractionResult
from docmind.extraction.word import extract_docx


def _build_docx_with_headings() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph(
        "This is the introduction body text explaining the topic in detail."
    )
    doc.add_paragraph("Background", style="Heading 1")
    doc.add_paragraph(
        "This is the background section with different supporting details."
    )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_structural_uses_real_docx_headings_as_section_boundaries() -> None:
    extraction = extract_docx(_build_docx_with_headings())
    assert len(extraction.headings) == 2  # sanity check on Week 9-10's extractor

    # Small enough budget that the two sections can't merge into one chunk,
    # large enough that neither section alone needs splitting further.
    chunks = chunk_structural(extraction, target_tokens=25)

    assert len(chunks) == 2
    assert chunks[0].text.startswith("Introduction")
    assert chunks[1].text.startswith("Background")


def test_structural_keeps_preamble_before_first_heading() -> None:
    doc = DocxDocument()
    doc.add_paragraph("This intro text appears before any heading.")
    doc.add_paragraph("Findings", style="Heading 1")
    doc.add_paragraph("The findings section body text.")
    buffer = BytesIO()
    doc.save(buffer)

    extraction = extract_docx(buffer.getvalue())
    chunks = chunk_structural(extraction, target_tokens=5)  # forces a split

    assert len(chunks) == 2
    assert chunks[0].text.startswith("This intro text")
    assert chunks[1].text.startswith("Findings")


def test_structural_falls_back_to_paragraphs_without_headings() -> None:
    text = "First paragraph here.\n\nSecond paragraph here.\n\nThird paragraph here."
    extraction = ExtractionResult(
        text=text, pages=[ExtractedPage(page_number=1, text=text)]
    )

    chunks = chunk_structural(extraction, target_tokens=1000)

    assert len(chunks) == 1  # generous budget groups every paragraph together
    assert "First paragraph" in chunks[0].text
    assert "Third paragraph" in chunks[0].text


def test_structural_splits_paragraphs_across_budget() -> None:
    text = "\n\n".join(
        f"Paragraph number {i} with some filler words in it." for i in range(10)
    )
    extraction = ExtractionResult(
        text=text, pages=[ExtractedPage(page_number=1, text=text)]
    )

    chunks = chunk_structural(extraction, target_tokens=20)  # small budget

    assert len(chunks) > 1


def test_structural_oversized_paragraph_becomes_its_own_chunk() -> None:
    huge_paragraph = "word " * 500  # one paragraph, no internal \n\n to split on

    extraction = ExtractionResult(
        text=huge_paragraph, pages=[ExtractedPage(page_number=1, text=huge_paragraph)]
    )

    chunks = chunk_structural(extraction, target_tokens=10)  # tiny budget

    assert len(chunks) == 1
    assert chunks[0].text.strip() == huge_paragraph.strip()


def test_structural_never_merges_across_pages() -> None:
    extraction = ExtractionResult(
        text="",
        pages=[
            ExtractedPage(page_number=1, text="Short page one text."),
            ExtractedPage(page_number=2, text="Short page two text."),
        ],
    )

    chunks = chunk_structural(extraction, target_tokens=1000)

    assert len(chunks) == 2
    assert chunks[0].page_numbers == [1]
    assert chunks[1].page_numbers == [2]


def test_structural_splits_windows_line_endings_into_paragraphs() -> None:
    # Regression test: a Windows-authored .txt file has literal "\r\n\r\n"
    # between paragraphs. Splitting on "\n\n" alone misses this entirely
    # and treats the whole file as one paragraph (found via Docs/DocMind.txt).
    text = "First paragraph.\r\n\r\nSecond paragraph.\r\n\r\nThird paragraph."
    extraction = ExtractionResult(
        text=text, pages=[ExtractedPage(page_number=1, text=text)]
    )

    chunks = chunk_structural(extraction, target_tokens=5)  # tiny budget

    assert len(chunks) == 3


def test_structural_empty_extraction_returns_no_chunks() -> None:
    extraction = ExtractionResult(text="", pages=[])
    assert chunk_structural(extraction) == []
